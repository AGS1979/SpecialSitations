import os
import re
import streamlit as st
from typing import List, Dict
import pdfplumber
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import tempfile
import base64
import yfinance as yf
from typing import List, Dict, Tuple

# --- Must be the first st.* command ---
st.set_page_config(page_title="Pre-IPO Memo Generator", layout="wide")

def get_base64_logo(path="logo.png"):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# --- Custom Header ---
logo_base64 = get_base64_logo()

st.markdown(f"""
    <style>
        /* 1. HIDE THE DEFAULT STREAMLIT HEADER */
        header {{
            visibility: hidden !important;
            height: 0 !important;
        }}

        /* 2. CREATE A CUSTOM FIXED HEADER */
        .custom-header {{
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            height: 4rem; /* Header height */
            background-color: #FFFFFF;
            display: flex;
            align-items: center;
            padding: 0 2rem;
            border-bottom: 1px solid #e0e0e0;
            z-index: 1000;
        }}
        
        .custom-header img {{
            height: 2rem; /* Logo height */
            width: auto;
        }}

        /* 3. PUSH THE MAIN CONTENT DOWN */
        .stApp {{
            margin-top: 4.5rem; /* Must be slightly larger than header height */
        }}
    </style>
    
    <div class="custom-header">
        <img src="data:image/png;base64,{logo_base64}" />
    </div>
""", unsafe_allow_html=True)


# --- Your App's Content Starts Here ---

st.title("Pre-IPO Investment Memo Generator")
st.write("Upload an IPO/DRHP PDF to generate a structured investment memo with optional Q&A.")

# ========== CONFIG ==========
try:
    # Securely fetch the API key from Streamlit's secrets
    DEEPSEEK_API_KEY = st.secrets["deepseek"]["api_key"]
except (KeyError, FileNotFoundError):
    st.error("DeepSeek API key not found. Please add it to your Streamlit secrets.")
    st.stop()

DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"


try:
    FMP_API_KEY = st.secrets["fmp"]["api_key"]
except (KeyError, FileNotFoundError):
    st.error("FMP API key not found in secrets. Please add it under [fmp][api_key].")
    st.stop()


# ==========================
# Report & Infographic Structures
# ==========================
REPORT_TEMPLATES = {
"Spin-Off or Split-Up": """
Transaction Overview
ParentCo and SpinCo details
Rationale (regulatory, strategic unlock, valuation arbitrage)
Distribution terms (ratio, eligibility, tax treatment)
ParentCo Post-Spin Outlook
Strategic focus
Financial profile and valuation
SpinCo Investment Case
Business model, growth drivers
Historical and pro forma financials
Independent valuation (e.g., Sum-of-the-Parts)
Valuation Analysis
Risks and Overhangs
Forced selling, low float, governance concerns
""",
"Mergers & Acquisitions": """
Deal Summary
Parties involved, consideration (cash/stock), premium
Regulatory/antitrust/board approval status
Target Company Analysis
Valuation vs. offer
Control premium vs. peers
Buyer’s Rationale and Financing
Strategic fit
Synergies and pro forma financials
Deal financing (debt, equity)
Shareholder Vote & Antitrust Risk
Key holders' stance
Timing and likelihood of deal closure
Spread Analysis and Arbitrage Opportunity
Deal spread
IRR scenarios based on timing/risk
""",
"Bankruptcy / Distressed / Restructuring": """
Situation Summary
Cause of distress
Filing date, jurisdiction, DIP terms
Capital Structure Analysis
Pre- and post-reorg structure
Seniority waterfall
Creditor classes and recovery potential
Valuation and Recovery Scenarios
Estimated Enterprise Value
Recovery per instrument (bonds, equity, unsecured)
Reorganization Plan and Exit Timeline
Conversion to equity, rights offering, warrants
Exit multiples
Catalysts and Legal Risks
Judge approval, creditor objections, asset sales
""",
"Activist Campaign": """
Activist Background
Fund profile, history, prior campaigns
Campaign Details
Demands (board seat, spin, buyback, etc.)
Timeline of engagement
Company's Response and Governance Profile
Management alignment, shareholder defense
Scenario Analysis
Status quo vs. activist success
Proxy fight implications
Valuation Impact
NPV of potential changes (e.g., spin-off value, ROIC uplift)
""",
"Regulatory or Legal Catalyst": """
Legal/Regulatory Background
Case/issue summary
Historical legal proceedings
Outcome Scenarios
Win, loss, settlement
Timeline
Financial and Strategic Implications
Fines, product approval, license loss
Revenue/EBITDA impact
Market Reaction History (if any)
Past similar cases
""",
"Asset Sales or Carve-Outs": """
Transaction Overview
Buyer, price, structure
Valuation vs. book and peers
Strategic Impact
Focus shift, deleveraging, margin profile
Use of Proceeds
Debt repayment, dividends, buybacks, capex
Re-rating Potential
EBITDA margin uplift, return metrics
""",
"Capital Raising or Buyback Catalyst": """
Transaction Mechanics
Size, dilution, instrument type
Capital Structure Post-Deal
Leverage ratios, interest burden
Shareholder Implications
Accretion/dilution
EPS impact
Buyback Analysis (if applicable)
Repurchase pace, valuation support
"""
}

FALLBACK_META = [
    ("💼", "border-blue-600", "bg-blue-50"),
    ("🏢", "border-sky-600", "bg-sky-50"),
    ("🌐", "border-indigo-600", "bg-indigo-50"),
    ("🧩", "border-purple-600", "bg-purple-50"),
    ("📊", "border-green-600", "bg-green-50"),
    ("📈", "border-emerald-600", "bg-emerald-50"),
    ("👥", "border-yellow-600", "bg-yellow-50"),
    ("⚠️", "border-red-600", "bg-red-50"),
    ("💡", "border-pink-600", "bg-pink-50"),
    ("🧠", "border-gray-600", "bg-gray-50"),
]



# ==========================
# Text Extractors
# ==========================

def extract_text_from_pdf(file):
    try:
        with pdfplumber.open(file) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    except Exception as e:
        return f"[ERROR extracting PDF: {e}]"

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"[ERROR extracting DOCX: {e}]"

# ==========================
# Memo Generation
# ==========================


def resolve_company_to_ticker(company_name: str) -> str:
    prompt = f"What is the stock ticker (FMP-compatible) for the public company '{company_name}'?"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0
    }

    try:
        res = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
        res.raise_for_status()
        ticker = res.json()["choices"][0]["message"]["content"].strip()
        return re.sub(r'[^A-Z\.]', '', ticker)  # sanitize
    except:
        return None

def get_ev_ebitda_multiple(ticker: str, fmp_key: str) -> float:
    url = f"https://financialmodelingprep.com/api/v3/key-metrics-ttm/{ticker}?apikey={fmp_key}"
    try:
        r = requests.get(url)
        data = r.json()
        if isinstance(data, list) and data:
            return float(data[0].get("enterpriseValueOverEBITDATTM", 0))
    except:
        return 0.0


def fetch_fundamentals_yf(ticker: str) -> Tuple[float, float, float]:
    """
    Returns (market_cap, net_debt, ttm_ebitda) via Yahoo Finance.
    """
    try:
        t = yf.Ticker(ticker)
        info = t.info or {}
        market_cap = info.get("marketCap", 0) or 0
        total_debt = info.get("totalDebt", 0) or 0
        # Some tickers put short‑term investments under a different key
        cash = info.get("cashAndShortTermInvestments", info.get("cash", 0)) or 0
        net_debt = total_debt - cash
        ebitda = info.get("ebitda", 0) or 0
        return float(market_cap), float(net_debt), float(ebitda)
    except Exception:
        return 0.0, 0.0, 0.0



def clean_markdown(text):
    text = re.sub(r'^[ \t\-]{3,}$', '', text, flags=re.MULTILINE)   # drop lines of --- or ***  
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'^- ', '• ', text, flags=re.MULTILINE)
    return text.strip()

def truncate_safely(text, limit=7000):
    return text[:limit]

def fetch_fundamentals_yf(ticker: str) -> Tuple[float, float, float]:
    """
    Returns (market_cap, net_debt, ttm_ebitda) via Yahoo Finance.
    """
    try:
        t = yf.Ticker(ticker)
        info = t.info or {}
        market_cap = info.get("marketCap", 0) or 0
        total_debt = info.get("totalDebt", 0) or 0
        cash = info.get("cashAndShortTermInvestments", info.get("cash", 0)) or 0
        net_debt = total_debt - cash
        ebitda = info.get("ebitda", 0) or 0
        return float(market_cap), float(net_debt), float(ebitda)
    except Exception:
        return 0.0, 0.0, 0.0


def generate_special_situation_note(
    company_name: str,
    situation_type: str,
    uploaded_files: list,
    valuation_mode: str = None,
    parent_peers: str = "",
    spinco_peers: str = "",
    fmp_key: str = ""  # no longer used but kept for signature compat
):
    # 1) Extract text
    combined_text = ""
    for file in uploaded_files:
        if file.name.endswith(".pdf"):
            combined_text += extract_text_from_pdf(file) + "\n"
        elif file.name.endswith(".docx"):
            combined_text += extract_text_from_docx(file) + "\n"
        else:
            combined_text += f"[Unsupported file: {file.name}]\n"

    # 2) Select template
    structure = REPORT_TEMPLATES.get(situation_type)
    if not structure:
        raise ValueError(f"Unsupported situation type: {situation_type}")

    # 3) Build valuation_section only for spin‑offs
    valuation_section = ""
    if situation_type == "Spin-Off or Split-Up" and valuation_mode:

        def process_peers(raw: str):
            names   = [n.strip() for n in raw.split(",") if n.strip()]
            tickers = [resolve_company_to_ticker(n) for n in names]
            mults   = [get_ev_ebitda_multiple(t, fmp_key) for t in tickers if t]
            avg     = round(sum(mults) / len(mults), 2) if mults else None
            return names, tickers, mults, avg

        if valuation_mode == "Let AI choose peers":
            # Ask DeepSeek for peers
            prompt = (
                f"List 5 large, publicly-traded companies most comparable to {company_name}, "
                "across its automation and aerospace segments, separated by commas."
            )
            resp = requests.post(
                DEEPSEEK_API_URL,
                headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
                json={"model":"deepseek-chat","messages":[{"role":"user","content":prompt}],"temperature":0}
            )
            resp.raise_for_status()
            body = resp.json().get("choices", [])
            ai_text = body[0].get("message",{}).get("content","") if body else ""
            peer_names = [n.strip() for n in ai_text.split(",") if n.strip()]
            peer_tickers = [resolve_company_to_ticker(n) for n in peer_names if resolve_company_to_ticker(n)]
            raw_mults    = [get_ev_ebitda_multiple(t, fmp_key) for t in peer_tickers]
            peer_mults   = [m for m in raw_mults if isinstance(m,(int,float))]
            avg_mult     = round(sum(peer_mults)/len(peer_mults),2) if peer_mults else None

            # Fetch fundamentals via yfinance
            ticker        = resolve_company_to_ticker(company_name)
            actual_mc, debt, ebitda = fetch_fundamentals_yf(ticker)
            ev_est        = (avg_mult or 0) * ebitda
            equity_est    = ev_est - debt
            upside_pct    = ((equity_est/actual_mc)-1)*100 if actual_mc else None

            valuation_section = f"""
# Valuation Analysis

**AI‑Selected Peers**: {', '.join(peer_names)}  
**Peer EV/EBITDA multiples**: {peer_mults}  
**Average EV/EBITDA**: {avg_mult or 'N/A'}  

**{company_name} TTM EBITDA**: ${ebitda:,.0f} mm  
**Estimated Enterprise Value**: {avg_mult or 0}×{ebitda:,.0f} = ${ev_est:,.0f} mm  
**Net Debt**: ${debt:,.0f} mm  
**Implied Equity Value**: ${equity_est:,.0f} mm  

**Actual Market Cap**: ${actual_mc:,.0f} mm  
**Implied Upside**: {f"{upside_pct:.1f}%" if upside_pct is not None else 'N/A'}  
"""
        elif valuation_mode == "I'll enter peer company names":
            p_names, p_tickers, p_mults, p_avg = process_peers(parent_peers)
            s_names, s_tickers, s_mults, s_avg = process_peers(spinco_peers)
            valuation_section = f"""
# Valuation Analysis

Peer companies provided by user:

**ParentCo Peers**: {', '.join(p_names)}  
EV/EBITDA multiples: {p_mults} (avg {p_avg or 'N/A'})  

**SpinCo Peers**: {', '.join(s_names)}  
EV/EBITDA multiples: {s_mults} (avg {s_avg or 'N/A'})  

Apply these averages to the TTM EBITDA from the docs for standalone valuations, then compare to ParentCo’s market cap for the unlock.
"""
        else:
            valuation_section = """
# Valuation Analysis

Identify relevant peers for ParentCo and SpinCo. Fetch their LTM EV/EBITDA multiples, compute averages, and multiply by TTM EBITDA for a SOTP comparison against ParentCo’s current market cap.
"""

    # 4) Assemble prompt
    prompt = f"""
You are an institutional investment analyst writing a professional memo on a special situation involving {company_name}.
The situation is: **{situation_type}**

Below is the internal company information extracted from various files:
\"\"\"{truncate_safely(combined_text)}\"\"\"

{valuation_section}

Using the structure below, generate a well-written investment memo. Be factual, insightful, and clear.
Structure:
{structure}
"""

    # 5) Call DeepSeek
    response = requests.post(
        DEEPSEEK_API_URL,
        headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
        json={"model":"deepseek-chat","messages":[{"role":"user","content":prompt}],"temperature":0.3}
    )
    response.raise_for_status()
    memo = clean_markdown(response.json()["choices"][0]["message"]["content"])

    # 6) Build and return .docx
    memo_dict = split_into_sections(memo, structure)
    doc = format_memo_docx(memo_dict, company_name, situation_type)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name




def split_into_sections(text: str, template: str) -> Dict[str, str]:
    sections = {}
    titles = [line.split('(')[0].strip() for line in template.strip().split('\n') if line.strip()]
    if not titles:
        return {"Memo": text.strip()}

    pattern = re.compile(r'^(' + '|'.join(map(re.escape, titles)) + r')\s*$', re.MULTILINE | re.IGNORECASE)
    matches = list(pattern.finditer(text))

    if not matches:
        return {"Memo": text.strip()}

    for i, match in enumerate(matches):
        title = match.group(1).strip()
        start_of_content = match.end()
        end_of_content = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start_of_content:end_of_content].strip()
        canonical_title = next((t for t in titles if t.lower() == title.lower()), title)
        sections[canonical_title] = content

    return sections

def format_memo_docx(memo_dict: dict, company_name: str, situation_type: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Aptos Display'
    style.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{company_name} – {situation_type} Investment Memo")
    title_run.font.name = 'Aptos Display'
    title_run.font.size = Pt(20)
    title_run.bold = True
    doc.add_paragraph()

    for section_title, content in memo_dict.items():
        heading = doc.add_paragraph()
        run = heading.add_run(section_title)
        run.bold = True
        run.font.size = Pt(14)
        heading.paragraph_format.space_after = Pt(6)    # heading‑to‑text gap
        for para in content.strip().split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6)    # tighten between paras
                p.paragraph_format.line_spacing = 1.3
        
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    return doc

# ==========================
# Infographic Generation
# ==========================

def extract_sections_from_docx_for_infographic(file, situation_type: str) -> Dict[str, str]:
    toc = REPORT_TEMPLATES.get(situation_type)
    if not toc:
        return {}
    
    expected_titles = {t.strip().lower() for t in toc.strip().splitlines() if t.strip()}
    doc = Document(file)
    sections = {}
    current_heading = None
    current_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        if text.lower() in expected_titles:
            if current_heading and current_text:
                sections[current_heading] = "\n".join(current_text).strip()
            current_heading = text
            current_text = []
        elif current_heading:
            current_text.append(text)

    if current_heading and current_text:
        sections[current_heading] = "\n".join(current_text).strip()
    
    return sections

def summarize_section_with_deepseek(section_title, section_text):
    prompt = f"""
You are an institutional research analyst preparing a financial infographic.
Summarize the section titled \"{section_title}\" into 3 to 5 concise bullet points.
Each point should be a single sentence, highlighting key insights clearly and professionally.
Section:
\"\"\"{section_text}\"\"\"
"""
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"].strip()

def build_infographic_html(company_name, sections):
    html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
    <title>{company_name} – Infographic</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
    <style>
        body {{ font-family: 'Inter', sans-serif; background-color: #f9fafb; color: #1f2937; }}
        .section-icon {{ font-size: 1.4rem; margin-right: 0.6rem; }}
    </style>
</head>
<body class="px-4 py-8 md:px-6 md:py-10 max-w-7xl mx-auto">
    <header class="text-center mb-12">
        <h1 class="text-3xl md:text-4xl font-bold text-gray-800 mb-2">{company_name} – Investment Memo Infographic</h1>
        <p class="text-sm text-gray-500">Generated by Aranca AI Platform</p>
    </header>
    <main class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">
"""
    with st.spinner("Summarizing sections for infographic..."):
        for idx, (title, section_text) in enumerate(sections.items()):
            icon, border_class, bg_class = FALLBACK_META[idx % len(FALLBACK_META)]
            try:
                summary = summarize_section_with_deepseek(title, section_text)
                cleaned_summary = summary.replace('**', '').replace('###', '').replace('##', '').replace('#', '')
                lines = [line.lstrip("•*- ").strip() for line in cleaned_summary.split("\n") if line.strip()]
                bullet_items = "\n".join(f"            <li>{line}</li>" for line in lines)
            except Exception as e:
                bullet_items = f"<li>Error generating summary: {e}</li>"
                st.warning(f"Could not summarize section: '{title}'")

            html += f"""
        <div class="shadow-lg rounded-xl p-5 transition-transform hover:scale-[1.02] duration-300 ease-in-out border-l-4 {border_class} {bg_class}">
            <h2 class="text-lg font-semibold text-gray-800 mb-3 flex items-center">
                <span class="section-icon">{icon}</span>{title}
            </h2>
            <ul class="list-disc text-sm text-gray-700 space-y-2 pl-5 leading-relaxed">
{bullet_items}
            </ul>
        </div>
"""
    html += """
    </main>
    <footer class="text-center mt-12">
        <p class="text-xs text-gray-400">This document is for informational purposes only. Not an investment advice.</p>
    </footer>
</body>
</html>
"""
    return html

# ==========================
# Streamlit App UI
# ==========================

st.title("📝 Special Situation Memo & Infographic Generator")
st.markdown("---")

st.sidebar.info("API key loaded from secrets.")

# --- Step 1: Memo Generation ---
st.header("Step 1: Generate Investment Memo")

company_name_memo = st.text_input("Enter Company Name", key="company_name_memo")
situation_type_memo = st.selectbox("Select Situation Type", options=list(REPORT_TEMPLATES.keys()), key="situation_type_memo")
valuation_mode = None   # ← make sure it always exists
parent_peers_raw = ""
spinco_peers_raw = ""


if situation_type_memo == "Spin-Off or Split-Up":
    st.markdown("### 🔍 Valuation Module (Optional)")

    valuation_mode = st.radio(
        "Do you want to provide peer companies for valuation, or let the model decide?",
        options=["Let AI choose peers", "I'll enter peer company names"],
        key="valuation_mode"
    )

    

    if valuation_mode == "I'll enter peer company names":
        parent_peers_raw = st.text_area("Enter ParentCo Peer Company Names (comma-separated)", key="parent_peers_raw")
        spinco_peers_raw = st.text_area("Enter SpinCo Peer Company Names (comma-separated)", key="spinco_peers_raw")

    else:
        st.info("AI will select peers using company descriptions and generate valuation logic automatically.")


uploaded_files_memo = st.file_uploader("Upload Public Documents (PDF, DOCX)", accept_multiple_files=True, key="uploaded_files_memo")

if st.button("Generate Memo"):
    if not company_name_memo or not situation_type_memo or not uploaded_files_memo:
        st.warning("Please fill in all fields and upload at least one document.")
    else:
        with st.spinner("Generating memo... This may take a moment."):
            try:
                memo_path = generate_special_situation_note(
                    company_name=company_name_memo,
                    situation_type=situation_type_memo,
                    uploaded_files=uploaded_files_memo,
                    valuation_mode=valuation_mode,
                    parent_peers = parent_peers_raw,
                    spinco_peers = spinco_peers_raw,
                    fmp_key=FMP_API_KEY
                )
                st.session_state.memo_path = memo_path
                st.session_state.company_name = company_name_memo
                st.session_state.situation_type = situation_type_memo
                
                st.success("Memo generated successfully!")
                with open(memo_path, "rb") as f:
                    st.download_button(
                        label="Download Memo (.docx)",
                        data=f,
                        file_name=f"{company_name_memo}_{situation_type_memo}_Memo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"An error occurred during memo generation: {e}")

st.markdown("\n\n---\n\n")

# --- Step 2: Infographic Generation ---
st.header("Step 2: Generate Infographic from Memo")
st.info("After generating the memo, upload the `.docx` file below to create an infographic.")

company_name_infographic = st.text_input("Confirm Company Name", value=st.session_state.get('company_name', ''), key="company_name_infographic")
situation_type_infographic = st.selectbox("Confirm Situation Type", options=list(REPORT_TEMPLATES.keys()), index=list(REPORT_TEMPLATES.keys()).index(st.session_state.get('situation_type')) if st.session_state.get('situation_type') else 0, key="situation_type_infographic")
uploaded_memo_infographic = st.file_uploader("Upload the generated Memo (.docx)", type=["docx"], key="uploaded_memo_infographic")


if st.button("Generate Infographic"):
    if not uploaded_memo_infographic or not company_name_infographic or not situation_type_infographic:
        st.warning("Please upload the memo and confirm the company name and situation type.")
    else:
        with st.spinner("Extracting sections..."):
            try:
                sections = extract_sections_from_docx_for_infographic(uploaded_memo_infographic, situation_type_infographic)
                if not sections:
                     st.error("Could not extract any sections from the document. Please ensure the memo contains headings matching the selected situation type.")
                else:
                    st.success(f"Successfully extracted {len(sections)} sections from the memo.")
                    html_content = build_infographic_html(company_name_infographic, sections)
                    
                    st.subheader("Infographic Preview")
                    st.components.v1.html(html_content, height=800, scrolling=True)

                    st.download_button(
                        label="Download Infographic (.html)",
                        data=html_content,
                        file_name=f"{company_name_infographic}_Infographic.html",
                        mime="text/html"
                    )
            except Exception as e:
                st.error(f"An error occurred during infographic generation: {e}")